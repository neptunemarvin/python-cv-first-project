from docx import Document
from docx.shared import Inches
import pyttsx3

def speak(text):
    pyttsx3.speak(text)

document = Document()

# adding image
document.add_picture(
    'flower.jpeg',
    width=Inches(1.75))

# input name & phone number & email
name = input('What is your name, yo?')
phone_number = input('Care to share your digits?')
email = input('How about that email address?')

document.add_paragraph(
    name + ' | ' + phone_number + ' | ' + email)

# about me
document.add_heading('About Me')

document.add_paragraph(
    input('Tell me about yourself ' + '\n'))

# work experience
document.add_heading('Work Experience')
p = document.add_paragraph()

company = input('Enter Company Name ')
from_date = input('From date ')
to_date = input('To date ')

p.add_run(company + ' ').bold = True
p.add_run(from_date + ' - ' + to_date + '\n').italic = True

# add position and or title
position_title = input(
    'Enter your position or title ')
p.add_run(position_title + ' ' + '\n').italic = True

experience_details = input(
    'Describe your experiance at ' + company + ' ')
p.add_run('   *  ' + experience_details).bulletpoints = True

# adding more experiences

while True:
    has_more_experiences = input(
        'Do you have more work experiences ? Yes or No ')
    if has_more_experiences.lower() == 'yes':
        p = document.add_paragraph()

        company = input('Enter Company Name ')
        from_date = input('From date ')
        to_date = input('To date ')

        p.add_run(company + ' ').bold = True
        p.add_run(from_date + ' - ' + to_date + '\n').italic = True

        # add position and or title
        position_title = input(
            'Enter your position or title ')
        p.add_run(position_title + ' ' + '\n')

        experience_details = input(
            'Describe your experiance at ' + company + ' ')
        p.add_run('   *  ' + experience_details).upper = True
    else:
        break

# enter skills
document.add_heading('Skills')
skill = input('Enter Skill ')
p = document.add_paragraph(skill)
# p.style = 'List bullet'

while True:
    has_more_skills = input('Do you have more skills? Yes or No ')
    if has_more_skills.lower() == 'yes':
        skill = input('Enter skill ')
        p = document.add_paragraph(skill)
        # p.style = 'List Bullet'
    else:
        break

#footer
section = document.sections[0]
footer = section.footer
p = footer.paragraphs[0]
p.text = "CV generated using python and mowing fool"

document.save('cv.docx')